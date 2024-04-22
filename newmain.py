import os
import openai
import docx
import fitz  # PyMuPDF

# Use an environment variable for the API key
# Replace 'YOUR_API_KEY' with your actual OpenAI API key
openai.api_key = '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'
#openai.api_key = os.getenv('OPENAI_API_KEY')

if not openai.api_key:
    print("OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
    exit(1)

def read_docx(file_path):
    doc = docx.Document(file_path)
    if doc.paragraphs:
        return doc.paragraphs[0].text
    return ''


def read_file(file_path):
    """Determines the file type and reads the file accordingly."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            print("Error: Unicode decode error for .txt file.")
            return None
    elif file_extension == '.pdf':
        try:
            text = ''
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            print(f"Error reading PDF file: {e}")
            return None
    else:
        print(f"Unsupported file type: {file_extension}")
        return None

def read_file_contents(file_path):
    """
    Reads the contents of a file and returns it as a string.
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def ask_openai_with_context(file_content, question):
    """
    Uses the OpenAI ChatCompletion to send a question with context provided from a file's content.
    Adjust 'model' as necessary.
    """
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",  # Adjust model as necessary
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": file_content},
            {"role": "user", "content": question}
        ]
    )
    return response.choices[0].message['content']

# Example usage  
if __name__ == "__main__":
    file_path = 'C:\\Users\\erez\\Downloads\\project plan.pdf'  # Update this with the actual path to your file
    file_content = read_file_contents(file_path)
    file_content = read_file(file_path)
    if file_content is not None:
        print(file_content)
    else:
        print("Failed to read the file.")
    
    question = "can you summary the content below?"  # Example question
    answer = ask_openai_with_context(file_content, question)
    print("Answer:", answer)
